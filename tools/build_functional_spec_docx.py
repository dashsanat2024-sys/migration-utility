#!/usr/bin/env python3
"""Build FUNCTIONAL_SPECIFICATION.docx with Word-friendly ASCII diagrams."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "FUNCTIONAL_SPECIFICATION.docx"

DIAGRAMS = {
    "actors": """
┌─────────────────────────────────────────────────────────────────────────┐
│                         PRIMARY ACTORS                                   │
│  Mapping Lead │ Business Analyst │ Product Owner │ Migration Engineer │ QA│
└───────────────────────────────┬─────────────────────────────────────────┘
                                │ uses browser
                                ▼
┌─────────────────────────────────────────────────────────────────────────┐
│                      MIGRATION UTILITY                                   │
│         React UI  ──►  FastAPI API  ──►  PostgreSQL                      │
└───────┬─────────────────────────────┬───────────────────────────────────┘
        │                             │
        │ CSV / JSON / XML extracts   │ Load payloads & export JSON
        ▼                             ▼
┌──────────────────┐         ┌────────────────────────────────────────┐
│ Legacy source    │         │ Destination APIs │ BI tools (Metabase) │
└──────────────────┘         └────────────────────────────────────────┘
""",
    "system_context": """
                    ┌──────────────────────┐
                    │  Migration team      │
                    │  (Web browser)       │
                    └──────────┬───────────┘
                               │ HTTPS
                               ▼
                    ┌──────────────────────┐
                    │  Migration Utility   │
                    │  (Platform)          │
                    └──┬───┬───┬───┬───┬───┘
           ┌───────────┘   │   │   │   └───────────┐
           ▼               ▼   ▼   ▼               ▼
    ┌────────────┐  ┌──────────┐  ┌────────┐  ┌──────────┐  ┌────────────┐
    │ PostgreSQL │  │ Landing  │  │ Kraken │  │   SAP    │  │ File export│
    │ metadata & │  │ zone     │  │ REST   │  │ ERP IDoc │  │ directory  │
    │ staging    │  │ files    │  │ API    │  │          │  │            │
    └────────────┘  └──────────┘  └────────┘  └──────────┘  └────────────┘
""",
    "container": """
┌─────────────┐
│ Web browser │
└──────┬──────┘
       │
       ├──────────────────────────────┐
       ▼                              ▼
┌─────────────────┐          ┌──────────────────────┐
│ Static SPA      │  /api/*  │ Serverless Python    │
│ React 19 + Vite │ ───────► │ FastAPI + Mangum     │
│ (Vercel CDN)    │          └──────────┬───────────┘
└─────────────────┘                     │
                          ┌─────────────┼─────────────┬─────────────┐
                          ▼             ▼             ▼             ▼
                   ┌───────────┐ ┌──────────┐ ┌───────────┐ ┌───────────┐
                   │ Neon      │ │ /tmp or  │ │ Kraken    │ │ SAP       │
                   │ PostgreSQL│ │ landing  │ │ adapter   │ │ adapter   │
                   └───────────┘ └──────────┘ └───────────┘ └───────────┘
                                                    │             │
                                               ┌────┴────┐   Mock / File export
                                               │  Mock   │
                                               └─────────┘
""",
    "local_dev": """
Developer machine
       │
       ├──► Vite :5174 ──proxy /api──► Uvicorn :8000 ──► Postgres :5433
       │
       └──► Nginx UI :3000 ──────────► Uvicorn :8000
""",
    "backend": """
API ROUTES (FastAPI)                DOMAIN SERVICES              ENGINES
─────────────────────               ───────────────              ───────
projects, workspace        ──►     FieldCatalogService          ValidationEngine
destination, fields        ──►     RuleSetService               TransformEngine
ingest                     ──►     IngestService                WorkflowEngine
rules, mapping             ──►     MappingService               MigrationPipeline
tariffs, selection         ──►     TariffService                SelectionEngine
runs, reconciliation       ──►     RunService / Reconciliation
                                           │
                                           ▼
                              PLUGIN / CONNECTOR REGISTRIES
                              DestinationPluginRegistry
                              ConnectorRegistry │ SchemaRegistry
                                           │
                                           ▼
                              DATA ACCESS: SQLAlchemy ORM + staging tables
""",
    "frontend": """
PAGES                    PROJECT SHELL                 WORKSPACE PANELS
───────                  ─────────────                 ────────────────
Dashboard ──► api.js     ProjectShell ──► Schema & Mapping (SchemaMappingScreen)
ProjectPage              (sidebar / mobile drawer) ──► Migration Wizard
                         │                      ──► Upload & Stage
                         │                      ──► Transform Rules
                         │                      ──► Tariff Mapping
                         │                      ──► Candidate Selection
                         │                      ──► Migration Runs
                         │                      ──► Reconciliation
                         └                      ──► Ingest Errors

Shared: SchemaMappingPanel, DestinationPluginCard, MigrationStepper
""",
    "plugins": """
REGISTERED PLUGINS                    PUBLISHED CONTRACT           UI
──────────────────                    ──────────────────           ───
kraken-billing-v3  ──┐
sap-crm-v1         ──┼──►  DestinationSchema  ──►  Mapping canvas
file-export-v1     ──┤     (fields, types,       (destination-first
mock-v1            ──┘      constraints)            rows)
""",
    "mapping_sequence": """
Step  Actor          Action
────  ─────          ──────
 1    User           Upload source extract CSV in Schema & Mapping UI
 2    UI → API       POST /api/projects/{id}/fields/{entity}/source
 3    API → DB       Parse file; store field catalog
 4    User           Click Auto-suggest mappings
 5    UI → API       POST /fields/{entity}/suggest-mappings
 6    API → UI       Return proposed source → target rows
 7    User           Edit transforms; apply to rule set
 8    UI → API       POST /fields/{entity}/apply-mappings/{rule_set_id}
 9    API → DB       Write field_mappings records
""",
    "ingest": """
File upload ──► Landing zone ──► Parser (CSV/JSON/XML) ──► Row validation
                                                                  │
                                    ┌─────────────────────────────┴────────────────────────────┐
                                    ▼ OK                                                         ▼ Fail
                            Staging table                                              ingest_errors
                         (per project + entity)                                      (reprocess queue)
""",
    "workflow": """
                    ┌─────────┐
                    │  START  │
                    └────┬────┘
                         ▼
                    ┌─────────┐     Mapping Lead / BA
              ┌────│  draft  │────────────────────────────┐
              │    └────┬────┘                            │
              │         │ Mapping Lead / BA               │
              │         ▼                                 │
              │    ┌───────────┐   BA / PO                │
              └───►│ in_review │──────────────► approved ─┼──► signed_off ──► END
                   └───────────┘                          │
                         ▲                                │
                         └──────── Product Owner ────────┘
""",
    "selection": """
Staged rows ────────┐
                    ├──► SelectionEngine ──► candidates ──► Migration run
Selection profile ──┘      (criteria + limits)
""",
    "pipeline": """
START ──► INGEST ──► VALIDATE ──► TRANSFORM ──► LOAD ──► Complete / Fail
            │            │             │           │
            └────────────┴─────────────┴───────────┴──► audit_logs
                                                      └──► load_records
""",
    "reconciliation": """
Staged count ────────┐
Validated count ─────┼──► Reconciliation funnel ──► Variance analysis
Transformed count ───┤                              └──► BI JSON export
Loaded count ────────┘
""",
    "er": """
PROJECT (1) ──────< (N) RULE_SET ──────< (N) VALIDATION_RULE
    │                      │
    │                      └──────< (N) FIELD_MAPPING
    │
    ├──────< (N) FIELD_CATALOG
    ├──────< (N) INGEST_FILE ──────< (N) INGEST_ERROR
    ├──────< (N) MIGRATION_RUN ─────< (N) BATCH
    │                      │
    │                      ├──────< (N) CANDIDATE
    │                      ├──────< (N) LOAD_RECORD
    │                      └──────< (N) AUDIT_LOG
    ├──────< (N) SELECTION_PROFILE ──< (N) SELECTION_CRITERION
    └──────< (N) TARIFF_MAPPING_SET ─< (N) TARIFF_MAPPING

RULE_SET ──────< (N) MAPPING_APPROVAL
""",
    "journey": """
Create project
      │
      ▼
Select destination plugin
      │
      ▼
Upload source extract ──► Auto-suggest mappings ──► Apply to rule set
      │
      ▼
Add validation rules ──► Workflow: draft → in_review → approved → signed_off
      │
      ├── [Utilities industry] ──► Configure tariff mapping
      │
      ▼
Upload & stage raw data
      │
      ▼
Configure candidate selection
      │
      ▼
Execute migration run
      │
      ▼
Reconciliation & export ──► Review ingest errors
""",
}


def set_cell_shading(cell, fill: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {qn("w:fill"): fill, qn("w:val"): "clear"})
    shading.append(shd)


def add_diagram(doc: Document, title: str, body: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    para = doc.add_paragraph()
    run = para.add_run(body.strip())
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8)
    para.paragraph_format.left_indent = Inches(0.15)
    para.paragraph_format.space_after = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(hdr[i], "E8EEF7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    title = doc.add_heading("Migration Utility — Functional Specification", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Product: Arthavi Migration Utility",
        "Version: 0.8.0",
        "Last updated: June 2026",
        "Live: https://migration-utility.vercel.app",
        "Repository: https://github.com/dashsanat2024-sys/migration-utility",
    ]:
        r = meta.add_run(line + "\n")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # 1 Purpose
    doc.add_heading("1. Purpose and scope", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "Migration Utility is a generic, industry-aware data migration platform that lets migration "
        "teams configure, validate, transform, and load structured records from a legacy source "
        "into a modern destination system, with full auditability and reconciliation."
    )
    p = doc.add_paragraph()
    p.add_run("Lifecycle: ").bold = True
    p.add_run("Extract → Validate → Transform → Load → Reconcile")
    doc.add_paragraph(
        "Configuration is stored in PostgreSQL and driven through a React web UI. The v0.8 flagship "
        "flow is destination-first schema mapping: the destination publishes its field contract via a "
        "destination plugin, and source extracts map into that contract."
    )

    doc.add_heading("1.2 In scope (v0.8)", level=2)
    add_table(
        doc,
        ["Area", "Capability"],
        [
            ["Project management", "Create projects by migration type, industry, integration approach"],
            ["Destination plugins", "Kraken Account (ST Water), SAP CRM, file export, mock"],
            ["Schema & mapping", "Upload source extract/catalog, auto-suggest mappings, apply to rule sets"],
            ["Ingest & staging", "CSV/JSON/XML upload, per-project staging tables, error queue"],
            ["Rules & transforms", "Validation rules, field mappings, transform types, workflow states"],
            ["Tariff mapping", "Utilities product/rate band → destination codes"],
            ["Candidate selection", "Filter staged records before run"],
            ["Migration runs", "Pipeline execution with batches, audit log, load records"],
            ["Reconciliation", "Funnel, variance, sample diffs, BI JSON export"],
            ["Deployment", "Vercel (UI + serverless API) + Neon PostgreSQL"],
        ],
    )

    doc.add_heading("1.3 Out of scope (v0.8)", level=2)
    add_table(
        doc,
        ["Item", "Status"],
        [
            ["User authentication / RBAC login", "Planned"],
            ["Live Kraken GraphQL schema introspection", "Planned"],
            ["Document migration, DB migration types", "Scaffolded (disabled in UI)"],
            ["Banking / healthcare industry templates", "Scaffolded (coming soon)"],
            ["Multi-tenant SaaS billing", "Not planned in v0.8"],
        ],
    )

    # 2 Actors
    doc.add_heading("2. Actors and stakeholders", level=1)
    add_diagram(doc, "Figure 2-1: Actors and system context", DIAGRAMS["actors"])
    add_table(
        doc,
        ["Actor", "Role", "Typical tasks"],
        [
            ["Mapping Lead", "Owns field mapping", "Upload extracts, map fields, submit for review"],
            ["Business Analyst", "Validates business rules", "Review mappings, approve rule sets"],
            ["Product Owner", "Sign-off authority", "Final sign-off before production runs"],
            ["Migration Engineer", "Executes runs", "Stage data, configure selection, run pipeline"],
            ["QA / Reconciliation", "Post-run verification", "Reconciliation funnel, variance, BI export"],
        ],
    )
    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run(
        "Workflow roles exist in the backend (mapping_lead, business_analyst, product_owner) but "
        "login/RBAC is not yet implemented — roles are passed via API request body in development."
    )

    # 3 System context
    doc.add_heading("3. System context (C4 — Level 1)", level=1)
    add_diagram(doc, "Figure 3-1: System context diagram", DIAGRAMS["system_context"])
    doc.add_paragraph(
        "System boundary: Migration Utility owns configuration, orchestration, staging, validation, "
        "transformation, load tracking, and reconciliation. It does not own the legacy source system "
        "or the destination platform of record."
    )

    # 4 Container
    doc.add_heading("4. Container architecture (C4 — Level 2)", level=1)
    add_diagram(doc, "Figure 4-1: Production container diagram (Vercel)", DIAGRAMS["container"])
    doc.add_heading("4.1 Container descriptions", level=2)
    add_table(
        doc,
        ["Container", "Technology", "Responsibility"],
        [
            ["React SPA", "Vite 6, React 19, React Router 7", "Dashboard, project workspace, mapping canvas, run monitoring"],
            ["FastAPI API", "Python 3.11+, Mangum on Vercel", "REST API, pipeline orchestration, plugin registry"],
            ["PostgreSQL", "Neon (prod), Docker (local)", "Projects, rules, catalogs, runs, staging tables, audit"],
            ["Landing zone", "Filesystem / /tmp", "Uploaded extract files before parse"],
            ["Target adapters", "Python connectors", "Mock, file export, Kraken, SAP load implementations"],
        ],
    )
    doc.add_heading("4.2 Local / Docker deployment", level=2)
    add_diagram(doc, "Figure 4-2: Local development topology", DIAGRAMS["local_dev"])

    # 5 Application architecture
    doc.add_heading("5. Application architecture (C4 — Level 3)", level=1)
    doc.add_heading("5.1 Backend component diagram", level=2)
    add_diagram(doc, "Figure 5-1: Backend layers", DIAGRAMS["backend"])
    doc.add_heading("5.2 Frontend component diagram", level=2)
    add_diagram(doc, "Figure 5-2: Frontend structure", DIAGRAMS["frontend"])
    doc.add_heading("5.3 URL and navigation model", level=2)
    add_table(
        doc,
        ["Route", "Screen", "Notes"],
        [
            ["/", "Dashboard", "Project list, create project wizard"],
            ["/projects/{slug}", "Project workspace", "Tab state in React (not in URL)"],
        ],
    )
    doc.add_paragraph("Default tab: Schema & Mapping. Mobile: hamburger drawer + sticky top bar.")

    doc.add_page_break()

    # 6 Functional domains
    doc.add_heading("6. Core functional domains", level=1)

    doc.add_heading("6.1 Project setup", level=2)
    doc.add_paragraph(
        "User creates a migration project with a profile that drives visible features and defaults."
    )
    doc.add_paragraph("Setup wizard steps:", style="List Number")
    for step in [
        "Migration type (Data / Document / DB — only Data enabled)",
        "Industry (Utilities, Generic; Banking/Healthcare scaffolded)",
        "Integration approach (API, File)",
        "Project details (name, slug, environment, connectors)",
    ]:
        doc.add_paragraph(step, style="List Number")
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        [
            ["PRJ-01", "System shall create a project with unique slug", "Must"],
            ["PRJ-02", "System shall resolve project by UUID or slug in API", "Must"],
            ["PRJ-03", "Profile shall gate tariff tab for non-utility industries", "Must"],
            ["PRJ-04", "System shall persist source/target connector keys on project", "Must"],
        ],
    )

    doc.add_heading("6.2 Destination plugin system", level=2)
    add_diagram(doc, "Figure 6-1: Destination plugin flow", DIAGRAMS["plugins"])
    add_table(
        doc,
        ["Plugin ID", "Label", "Adapter key", "Transport"],
        [
            ["kraken-billing-v3", "Kraken Account — Severn Trent Water", "kraken", "GraphQL · REST"],
            ["sap-crm-v1", "SAP Customer Master", "sap", "IDoc / BAPI"],
            ["file-export-v1", "JSON File Export", "file_export", "File system"],
            ["mock-v1", "Mock Destination", "mock", "In-memory"],
        ],
    )
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        [
            ["PLG-01", "System shall list available destination plugins", "Must"],
            ["PLG-02", "System shall return active plugin schema for a project", "Must"],
            ["PLG-03", "User shall swap destination plugin with orphan confirmation", "Must"],
            ["PLG-04", "User may upload custom destination schema CSV/JSON", "Should"],
            ["PLG-05", "Kraken plugin shall expose ST Water AccountType fields (~40+)", "Must"],
        ],
    )

    doc.add_heading("6.3 Schema and field mapping", level=2)
    add_diagram(doc, "Figure 6-2: Mapping workflow (sequence)", DIAGRAMS["mapping_sequence"])
    add_table(
        doc,
        ["Type", "Use case"],
        [
            ["copy", "Direct field copy"],
            ["constant", "Hardcoded value (e.g. migrationSource)"],
            ["lookup", "Enum translation (status, account type)"],
            ["conditional", "Y/N flags → boolean"],
            ["concat", "Join name parts"],
            ["date_format", "Date reformatting"],
        ],
    )
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        [
            ["MAP-01", "System shall render one row per destination schema field", "Must"],
            ["MAP-02", "System shall auto-suggest mappings with alias matching", "Must"],
            ["MAP-03", "Unmapped source columns shall appear as source-only rows", "Must"],
            ["MAP-04", "Mappings shall apply to a selected rule set", "Must"],
            ["MAP-05", "Mapping edits locked when rule set not in draft/in_review", "Must"],
        ],
    )

    doc.add_heading("6.4 Ingest and staging", level=2)
    add_diagram(doc, "Figure 6-3: Ingest pipeline", DIAGRAMS["ingest"])
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        [
            ["ING-01", "Support CSV, JSON, XML ingest", "Must"],
            ["ING-02", "Persist ingest file metadata and row counts", "Must"],
            ["ING-03", "Record failed rows with reason in ingest_errors", "Must"],
            ["ING-04", "Expose staging stats per entity", "Must"],
        ],
    )

    doc.add_heading("6.5 Rules, validation, and workflow", level=2)
    add_diagram(doc, "Figure 6-4: Rule set workflow states", DIAGRAMS["workflow"])
    doc.add_paragraph(
        "Validation rule types: required, format, in_list, range, cross_field, unique"
    )

    doc.add_heading("6.6 Tariff mapping (utilities)", level=2)
    doc.add_paragraph(
        "Map legacy product/rate codes to destination tariff/product codes. Tab visible when "
        "features.tariff_mapping is enabled (Utilities industry)."
    )

    doc.add_heading("6.7 Candidate selection", level=2)
    add_diagram(doc, "Figure 6-5: Selection engine", DIAGRAMS["selection"])

    doc.add_heading("6.8 Migration pipeline (runs)", level=2)
    add_diagram(doc, "Figure 6-6: Migration pipeline stages", DIAGRAMS["pipeline"])
    add_table(
        doc,
        ["Key", "Description"],
        [
            ["use_rules", "Apply approved rule set"],
            ["use_selection", "Filter by selection profile"],
            ["selection_profile_id", "Profile UUID"],
            ["require_approved_rules", "Block run if not approved"],
        ],
    )

    doc.add_heading("6.9 Reconciliation", level=2)
    add_diagram(doc, "Figure 6-7: Reconciliation funnel", DIAGRAMS["reconciliation"])

    doc.add_page_break()

    # 7 Data architecture
    doc.add_heading("7. Data architecture", level=1)
    doc.add_heading("7.1 Entity relationship (logical)", level=2)
    add_diagram(doc, "Figure 7-1: Entity relationships", DIAGRAMS["er"])
    doc.add_heading("7.2 Database phases (Alembic)", level=2)
    add_table(
        doc,
        ["Phase", "Key tables"],
        [
            ["0", "projects, migration_runs, batches, audit_logs"],
            ["1", "ingest_files, ingest_errors, dynamic staging"],
            ["2", "rule_sets, validation_rules, field_mappings"],
            ["3", "selection_profiles, selection_criteria, candidates"],
            ["4", "mapping_approvals, tariff_mapping_sets, tariff_mappings"],
            ["5", "load_records"],
            ["6", "field_catalogs"],
        ],
    )
    doc.add_heading("7.3 Workspace bootstrap", level=2)
    doc.add_paragraph("Single API call when opening a project:")
    add_diagram(
        doc,
        "API",
        "GET /api/projects/{slug}/workspace?entity=account\n\n"
        "Returns: project, plugin, destination_schema, catalog, rule_sets, entities",
    )

    # 8 Journey
    doc.add_heading("8. End-to-end user journey", level=1)
    add_diagram(doc, "Figure 8-1: End-to-end migration flow", DIAGRAMS["journey"])
    doc.add_heading("8.1 Recommended tab sequence", level=2)
    add_table(
        doc,
        ["Step", "UI tab", "Outcome"],
        [
            ["1", "Schema & Mapping", "Destination contract + field mappings"],
            ["2", "Upload & Stage", "Raw data in staging tables"],
            ["3", "Transform Rules", "Validation + workflow approval"],
            ["4", "Tariff Mapping", "Product translations (utilities)"],
            ["5", "Candidate Selection", "Record filter for run"],
            ["6", "Migration Runs", "Execute pipeline"],
            ["7", "Reconciliation", "Post-run verification"],
            ["8", "Ingest Errors", "Fix and reprocess failures"],
        ],
    )

    # 9 API
    doc.add_heading("9. API summary", level=1)
    doc.add_paragraph("Base path: /api — OpenAPI at /docs (local).")
    add_table(
        doc,
        ["Domain", "Prefix", "Key operations"],
        [
            ["Health", "/health, /health/live", "Liveness and DB check"],
            ["Projects", "/projects", "CRUD; slug or UUID lookup"],
            ["Workspace", "/projects/{ref}/workspace", "Bootstrap payload"],
            ["Destination", "/destination/plugins, /projects/{id}/destination/*", "Plugin schema and swap"],
            ["Fields", "/projects/{id}/fields/{entity}/*", "Catalog upload, suggest, apply"],
            ["Ingest", "/projects/{id}/ingest/*", "Upload, stats, errors"],
            ["Rules", "/projects/{id}/rules/*", "Rule sets, validation, workflow"],
            ["Runs", "/projects/{id}/runs, /runs/{id}/*", "Execute, audit, loads"],
            ["Reconciliation", "/projects/{id}/reconciliation/*", "Summary, export"],
        ],
    )

    # 10 NFR
    doc.add_heading("10. Non-functional requirements", level=1)
    add_table(
        doc,
        ["Category", "Requirement", "Current state"],
        [
            ["Availability", "Vercel auto-deploy from main", "Implemented"],
            ["Performance", "Workspace bootstrap ≤ 1 round-trip", "Implemented v0.8"],
            ["Performance", "Serverless cold start 3–5s", "Known limitation"],
            ["Security", "Authentication / RBAC", "Implemented v0.9 (AUTH_ENABLED)"],
            ["Enterprise", "Async worker + progress/resume", "Implemented v0.9"],
            ["Enterprise", "Data profiling + exception queue", "Implemented v0.9"],
            ["Mobile", "Responsive drawer navigation", "Implemented v0.8"],
            ["Testing", "60 automated backend tests", "pytest"],
        ],
    )

    # 11 Env
    doc.add_heading("11. Environment configuration", level=1)
    add_table(
        doc,
        ["Variable", "Purpose"],
        [
            ["DATABASE_URL", "PostgreSQL connection (Neon in prod)"],
            ["LANDING_ZONE_PATH", "Upload storage (/tmp on Vercel)"],
            ["CORS_ORIGINS", "Allowed frontend origins"],
            ["KRAKEN_MOCK_MODE", "Mock vs live Kraken load (default true)"],
            ["SAP_MOCK_MODE", "Mock vs live SAP load"],
            ["AUTH_ENABLED", "Enable JWT login and RBAC"],
            ["RUNNER_MODE", "api (sync) or worker (async queue)"],
            ["HTTP_PROXY / HTTPS_PROXY", "Corporate egress proxy"],
            ["CLIENT_CERT_PATH", "mTLS client certificate for destination APIs"],
        ],
    )

    # 15 Customer FAQ / RFP
    doc.add_heading("15. Customer FAQ / RFP summary", level=1)
    doc.add_paragraph(
        "Detailed procurement and security answers live in docs/CUSTOMER_FAQ_RFP.md. "
        "The one-page capability matrix for sales is docs/CAPABILITY_MATRIX.md."
    )
    add_table(
        doc,
        ["Question", "Answer (summary)"],
        [
            ["Vendor neutral?", "Yes — destination plugins; UI avoids lock-in"],
            ["On-prem / VPC?", "Yes — Docker Compose (API, UI, worker, Postgres)"],
            ["Proxy / mTLS?", "Yes — env-configured HTTP client for live loads"],
            ["Async runs + resume?", "Yes — worker polls queued runs; checkpoint resume"],
            ["Data profiling?", "Yes — on ingest; anomalies in UI and API"],
            ["RBAC / approvals?", "Partial — JWT when AUTH_ENABLED; no SSO yet"],
            ["Exception queue (HITL)?", "Yes — assign, override, resolve, audit"],
            ["100M+ row scale?", "Partial — chunked batches; customer sizing required"],
            ["SSO?", "Roadmap — not in v0.9"],
        ],
    )
    doc.add_paragraph("Runner deployment: docs/DEPLOYMENT_RUNNER.md")

    # 12 Use case
    doc.add_heading("12. Primary use case — Utilities / Kraken", level=1)
    add_table(
        doc,
        ["Legacy (source)", "Kraken (destination)", "Transform"],
        [
            ["CUST_ACCOUNT_NO", "number", "copy"],
            ["LEGACY_SYS_REF", "urn", "copy"],
            ["CUST_TYPE_FLAG", "accountType", "lookup D→DOMESTIC"],
            ["ACCT_STATUS_CODE", "status", "lookup A→ACTIVE"],
            ["STEPPED_RATE_FLAG", "isOnSteppedTariff", "conditional Y/N"],
            ["(constant)", "migrationSource", "constant TARGET_CMP"],
        ],
    )

    # 13 Roadmap
    doc.add_heading("13. Roadmap (functional)", level=1)
    add_table(
        doc,
        ["Feature", "Target"],
        [
            ["SSO / OIDC integration", "v1.0"],
            ["Live destination schema introspection", "v0.9+"],
            ["Document and DB migration types", "v1.0"],
            ["Banking / healthcare industry packs", "v1.0"],
        ],
    )

    # 14 Glossary
    doc.add_heading("14. Glossary", level=1)
    add_table(
        doc,
        ["Term", "Definition"],
        [
            ["Source", "Legacy system providing extract data"],
            ["Destination", "Target platform; publishes schema via plugin"],
            ["Destination plugin", "Module owning get_schema() contract"],
            ["Rule set", "Versioned validation rules + field mappings"],
            ["Staging table", "Per-project PostgreSQL table for extract rows"],
            ["Load record", "Persisted adapter request/response for one entity"],
            ["Reconciliation", "Post-run funnel, variance, and BI export"],
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph("For setup and deployment details see docs/PROJECT_DOCUMENTATION.md")
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
